"""
Tailor a base_resume.docx (converted from your PDF) to a job description using OpenAI or Hugging Face (fallback).

Output:
  - output/resumes/Anuraj_<CompanyName>.docx
  - output/resumes/Anuraj_<CompanyName>.pdf
  - output/descriptions/<CompanyName>_description.txt
"""

import os
import json
import re
import subprocess
import requests
from datetime import datetime
from docx import Document
from openai import OpenAI

# --- Load API keys
OPENAI_KEY = os.getenv("OPENAI_API_KEY")
HF_KEY = os.getenv("HUGGINGFACE_API_KEY")

client = None
if OPENAI_KEY:
    try:
        client = OpenAI(api_key=OPENAI_KEY)
    except Exception as e:
        print(f"⚠️ OpenAI init failed: {e}")
        client = None

BASE_DOCX = "data/base_resume.docx"
OUT_DIR_RESUMES = "output/resumes"
OUT_DIR_DESC = "output/descriptions"

# --- Utility
def sanitize_name(name: str) -> str:
    s = re.sub(r"[^\w\-]+", "_", name.strip())
    s = re.sub(r"__+", "_", s)
    return s.strip("_")[:120]

def extract_docx_text(path=BASE_DOCX) -> str:
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(texts)

# --- Prompt Template
PROMPT_TEMPLATE = """
You are an expert QA Automation resume writer and ATS optimizer.

Given:
1) Base resume text
2) Job title and job description snippet

Task:
- Produce JSON with keys: "summary", "skills" (array), "experience_updates" (array)
- Use relevant QA phrases: Selenium, BDD, Cucumber, Java, REST Assured, Jenkins, CI/CD.
- Keep it short, accurate, and ATS-optimized.

Return valid JSON.

Base Resume:
\"\"\"{base_text}\"\"\"

Job Title: {job_title}

Job Description:
\"\"\"{job_desc}\"\"\"
"""

# --- Hugging Face Fallback
def hf_tailor_request(base_text, job_title, job_desc):
    """Use Hugging Face Inference API as free fallback."""
    if not HF_KEY:
        raise RuntimeError("Hugging Face API key not set in .env")
    # smaller free model for fallback
    url = "https://api-inference.huggingface.co/models/HuggingFaceH4/zephyr-7b-beta"
    headers = {"Authorization": f"Bearer {HF_KEY}"}
    prompt = PROMPT_TEMPLATE.format(base_text=base_text, job_title=job_title, job_desc=job_desc)
    payload = {"inputs": prompt, "parameters": {"max_new_tokens": 512, "temperature": 0.4}}
    r = requests.post(url, headers=headers, json=payload, timeout=90)
    r.raise_for_status()
    data = r.json()
    if isinstance(data, list) and len(data) and "generated_text" in data[0]:
        content = data[0]["generated_text"]
    elif isinstance(data, dict) and "generated_text" in data:
        content = data["generated_text"]
    else:
        content = str(data)
    try:
        return json.loads(re.search(r"(\{.*\})", content, flags=re.S).group(1))
    except Exception:
        return {"summary": "Tailored summary (HF fallback).",
                "skills": ["QA", "Automation", "Selenium", "Java"],
                "experience_updates": ["Adapted resume via free model."]}

# --- OpenAI Primary
def request_tailored_sections(base_text: str, job_title: str, job_desc: str) -> dict:
    prompt = PROMPT_TEMPLATE.format(base_text=base_text, job_title=job_title, job_desc=job_desc)
    # Try OpenAI first if client exists
    if client:
        try:
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=700
            )
            content = resp.choices[0].message.content
            return json.loads(re.search(r"(\{.*\})", content, flags=re.S).group(1))
        except Exception as e:
            err = str(e)
            if "quota" in err.lower() or "rate" in err.lower():
                print("⚠️ OpenAI quota exceeded — using Hugging Face fallback.")
                return hf_tailor_request(base_text, job_title, job_desc)
            print(f"⚠️ OpenAI error ({e}), using HF fallback.")
            return hf_tailor_request(base_text, job_title, job_desc)
    else:
        print("⚠️ No OpenAI key available — using Hugging Face fallback.")
        return hf_tailor_request(base_text, job_title, job_desc)

# --- Apply updates
def apply_updates_to_docx(out_docx_path: str, updates: dict):
    doc = Document(BASE_DOCX)
    headings = {
        "summary": ["profile summary", "profile", "summary"],
        "skills": ["core competencies", "it skills", "skills", "core skills"],
        "experience": ["work experience", "experience", "projects"]
    }
    para_texts = [p.text.strip() for p in doc.paragraphs]
    lower_texts = [t.lower() for t in para_texts]

    def replace_section_by_heading(heading_list, new_lines):
        for h in heading_list:
            for idx, t in enumerate(lower_texts):
                if h in t:
                    for i, line in enumerate(new_lines):
                        if idx + 1 + i < len(doc.paragraphs):
                            doc.paragraphs[idx + 1 + i].text = line
                        else:
                            doc.add_paragraph(line)
                    return True
        return False

    if updates.get("summary"):
        lines = [l.strip() for l in updates["summary"].split("\n") if l.strip()]
        replace_section_by_heading(headings["summary"], lines)

    if updates.get("skills"):
        replace_section_by_heading(headings["skills"], [", ".join(updates["skills"])])

    if updates.get("experience_updates"):
        for b in updates["experience_updates"]:
            doc.add_paragraph("- " + b)

    os.makedirs(os.path.dirname(out_docx_path), exist_ok=True)
    doc.save(out_docx_path)
    print(f"✅ Saved tailored DOCX: {out_docx_path}")

def convert_docx_to_pdf(docx_path):
    if not os.path.exists(docx_path):
        raise FileNotFoundError(docx_path)
    pdf_path = docx_path.rsplit(".", 1)[0] + ".pdf"
    cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", os.path.dirname(docx_path), docx_path]
    subprocess.run(cmd, check=False)
    if os.path.exists(pdf_path):
        print(f"✅ Saved PDF: {pdf_path}")
    return pdf_path

def tailor_and_save(job_title: str, company_name: str, job_desc: str):
    base_text = extract_docx_text(BASE_DOCX)
    updates = request_tailored_sections(base_text, job_title, job_desc)
    company_clean = sanitize_name(company_name or job_title)
    fname = f"Anuraj_{company_clean}"
    out_docx = os.path.join(OUT_DIR_RESUMES, fname + ".docx")
    out_desc = os.path.join(OUT_DIR_DESC, f"{company_clean}_description.txt")

    os.makedirs(OUT_DIR_RESUMES, exist_ok=True)
    os.makedirs(OUT_DIR_DESC, exist_ok=True)

    with open(out_desc, "w", encoding="utf-8") as f:
        f.write(job_desc)

    apply_updates_to_docx(out_docx, updates)
    try:
        out_pdf = convert_docx_to_pdf(out_docx)
    except Exception as e:
        print(f"⚠️ PDF conversion failed: {e}")
        out_pdf = None
    return out_docx, out_pdf

if __name__ == "__main__":
    print("🚀 Tailoring resume test...")
    docx_path, pdf_path = tailor_and_save("QA Automation Engineer", "TechNova", "QA automation, Selenium, Java, API testing.")
    print("✅ Done:", docx_path, pdf_path)
