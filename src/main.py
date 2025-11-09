# src/main.py
"""
Main orchestrator for QA Job Apply Agent.
- Converts PDF -> DOCX if needed
- Searches jobs
- For each job: tries OpenAI tailoring, falls back to local heuristic tailoring
- Saves outputs and notifies via Telegram
"""
import os
import traceback
from scripts.convert_pdf_to_docx import convert as convert_pdf_to_docx
from src.job_search import search_jobs
from src.resume_tailor import tailor_and_save  # uses OpenAI
from src.telegram_bot import send_message
from dotenv import load_dotenv
import time

load_dotenv()  # load .env for local runs

# --- Local fallback tailoring (simple, free heuristic)
def local_tailor_and_save(job_title: str, company_name: str, job_desc: str):
    """
    Simple free fallback:
    - Reads base_resume.docx text
    - Finds top keywords from job_desc
    - Inserts keywords into top-of-resume summary and skills section
    - Saves as Anuraj_<Company>_local.docx (no PDF conversion)
    """
    from docx import Document
    import re, json

    BASE_DOCX = "data/base_resume.docx"
    OUT_DIR_RESUMES = "output/resumes"
    OUT_DIR_DESC = "output/descriptions"
    os.makedirs(OUT_DIR_RESUMES, exist_ok=True)
    os.makedirs(OUT_DIR_DESC, exist_ok=True)

    # read docx
    doc = Document(BASE_DOCX)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text and not p.text.isspace()])

    # extract keywords (very simple): words that appear frequently & capitalized tech terms
    job_text = job_desc.lower()
    keywords = []
    common_terms = ["selenium", "java", "python", "bdd", "cucumber", "rest", "api", "ci/cd", "jenkins", "page object model", "webdriver"]
    for t in common_terms:
        if t in job_text:
            keywords.append(t.upper() if "/" not in t else t.upper())

    # fallback: get top nouns (very naive)
    tokens = [w.strip(".,()") for w in job_text.split() if len(w) > 3]
    freq = {}
    for w in tokens:
        freq[w] = freq.get(w, 0) + 1
    # get top 5 frequent tokens (excluding stop-ish words)
    stop = set(["with","that","which","using","experience","years","required","role","will","work","and","the","for","in","a","to"])
    top_tokens = [k for k, v in sorted(freq.items(), key=lambda x: x[1], reverse=True) if k not in stop][:5]
    for t in top_tokens:
        if t.lower() not in [k.lower() for k in keywords]:
            keywords.append(t.upper())

    # Insert keywords into the first "summary" paragraph or top area
    if len(doc.paragraphs) > 0:
        # prepend a short tailored summary line
        summary_line = f"Tailored for {job_title} — highlights: {', '.join(keywords)}"
        doc.paragraphs[0].text = summary_line + "\n" + doc.paragraphs[0].text

    # Try to find a "Skills" heading and insert keywords line
    inserted_skills = False
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip().lower()
        if "skills" in txt or "core competencies" in txt or "it skills" in txt:
            # replace next paragraph with keywords
            if i+1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = ", ".join([k.upper() for k in keywords])
            else:
                doc.add_paragraph(", ".join([k.upper() for k in keywords]))
            inserted_skills = True
            break
    if not inserted_skills:
        # append skills at the end
        doc.add_paragraph("Tailored Skills: " + ", ".join([k.upper() for k in keywords]))

    # Save docx
    company_clean = re.sub(r"[^\w]+", "_", (company_name or job_title).strip())[:120]
    fname = f"Anuraj_{company_clean}_LOCAL.docx"
    out_docx = os.path.join(OUT_DIR_RESUMES, fname)
    doc.save(out_docx)

    # save job desc
    out_desc = os.path.join(OUT_DIR_DESC, f"{company_clean}_description.txt")
    with open(out_desc, "w", encoding="utf-8") as f:
        f.write(job_desc)

    return out_docx, None  # no PDF conversion for local fallback

def main():
    try:
        send_message("🚀 Starting QA Job Apply Agent...")
        # Step 1: Ensure resume converted
        if not os.path.exists("data/base_resume.docx") and os.path.exists("data/Resume-ANURAJ.pdf"):
            send_message("📄 Converting PDF → DOCX resume...")
            convert_pdf_to_docx()
            time.sleep(1)
        else:
            send_message("✅ Base resume ready.")

        # Step 2: Job Search
        send_message("🔍 Searching jobs (Bangalore / Remote)...")
        jobs = search_jobs(query="QA Automation Engineer", location_keywords=["Bangalore", "Bengaluru", "Remote"], max_results=10)
        if not jobs:
            send_message("❌ No jobs found this run.")
            return
        send_message(f"✅ Found {len(jobs)} jobs. Tailoring top 3...")

        # Step 3: Tailor resumes for top jobs
        for i, job in enumerate(jobs[:3]):
            title = job.get("title", "QA Engineer")
            snippet = job.get("snippet", "")
            company_name = job.get("title", "").split("-")[0].strip() or "Company"
            send_message(f"✂️ Tailoring resume for *{title}*...")

            # Try OpenAI-based tailoring first
            try:
                docx_path, pdf_path = tailor_and_save(title, company_name, snippet)
                send_message(f"📄 Tailored (OpenAI) ready:\n{docx_path}\n{pdf_path or 'PDF conversion skipped'}\n🔗 {job.get('url')}")
            except Exception as e:
                err = str(e)
                # Recognize quota / rate-limit message
                if "insufficient_quota" in err or "RateLimitError" in err or "quota" in err.lower():
                    send_message("⚠️ OpenAI quota exhausted or rate-limited. Falling back to local heuristic tailoring (free).")
                    docx_path, pdf_path = local_tailor_and_save(title, company_name, snippet)
                    send_message(f"📄 Tailored (LOCAL) ready:\n{docx_path}\n(Upload to GitHub Actions artifacts after run)\n🔗 {job.get('url')}")
                else:
                    # For other OpenAI errors, still fallback instead of crashing
                    send_message(f"⚠️ OpenAI error: {err[:200]}\nFalling back to local tailoring.")
                    docx_path, pdf_path = local_tailor_and_save(title, company_name, snippet)
                    send_message(f"📄 Tailored (LOCAL) ready:\n{docx_path}\n🔗 {job.get('url')}")

        send_message("✅ All done! Check output/resumes for files.")
    except Exception as e:
        tb = traceback.format_exc()
        msg = f"❌ Agent crashed with error:\n{e}\n\nTraceback:\n{tb[:3000]}"
        print(msg)
        try:
            send_message(msg)
        except:
            pass

if __name__ == "__main__":
    main()
