import openai, os
from docx import Document

def tailor_resume(template_path, job_description, output_path):
    doc = Document(template_path)
    resume_text = "\n".join([p.text for p in doc.paragraphs])
    
    prompt = f"""Optimize this resume text for the following job description so it scores above 90 ATS:
Resume:
{resume_text}
Job Description:
{job_description}
"""
    openai.api_key = os.getenv("OPENAI_API_KEY")
    completion = openai.ChatCompletion.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    new_text = completion.choices[0].message['content']
    new_doc = Document()
    new_doc.add_paragraph(new_text)
    new_doc.save(output_path)
